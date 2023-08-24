import os, global_var, logging
from flask import Flask, render_template, request, send_file, redirect
from werkzeug.exceptions import RequestEntityTooLarge
from functions import run_editor
from dotenv import load_dotenv, find_dotenv
import docx

# from flask_socketio import SocketIO
# this will come into use when we start using web sockets in order to get a better progress page running

load_dotenv(find_dotenv())
logging.basicConfig(level=logging.INFO, filename="log.log", filemode="w")
logger = logging.getLogger(__name__)

# Turn on debug level log statements from all libraries
logging.getLogger().setLevel(logging.DEBUG)

app = Flask(__name__)
app.config["UPLOAD_DIRECTORY"] = 'text_files/'
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024 #16MB
app.config["ALLOWED_EXTENSIONS"] = [".txt", ".docx"] #Would like to add LaTex and RTF compatibility later.
# socketio = SocketIO(app)

@app.route('/', methods=["GET", "POST"])
def index():
    if request.method == "GET":
        return render_template("index.html")


@app.route('/upload', methods=["POST"])
def upload():
    global_var.submit_text = "" #clear this variable in case the user clicked the back button
    global_var.key = request.form['key']
    if request.form['upload'] == "upload_file":
        try:
            file = request.files['file']
            if not file:
                return "Must upload a file"
            # Would be nice to check the file size here. Not sure if MAX_CONTENT checks it here or after file.save, 
            # but it takes a long time for it to check files that are 1 GB+
            extension = os.path.splitext(file.filename)[1]
            if extension not in app.config["ALLOWED_EXTENSIONS"]:
                logger.error("Unallowed extension uploaded")
                return "Cannot upload that file type. Must be '.txt' or '.docx'"
        except RequestEntityTooLarge:
            return "File is too large."

        if extension == ".txt":
            for paragraph in file.read().decode('utf-8', errors='ignore').split("\n"):
                global_var.submit_text += paragraph + "\n"

        if extension == ".docx":
            file = docx.Document(file)
            for paragraph in file.paragraphs:
                global_var.submit_text += paragraph.text + "\n"

    if request.form['upload'] == "upload_text":        
        #reset paragraph formatting sent by the HTML form
        for paragraph in request.form['text_box'].split("\r\n"):
            global_var.submit_text += paragraph + "\n"
        if global_var.submit_text == [""]:
            return "Text box is blank"
    return redirect('/progress')


@app.route('/progress', methods=["GET", "POST"])
def progress():
    chunk_count = (len(global_var.submit_text) // 4000) + 1
    if request.method == "GET": 
        return render_template("progress.html", chunks=chunk_count, wait=chunk_count * 15)
    if request.method == "POST":
        run_editor(global_var.submit_text, chunk_count)
        return redirect('/results')


@app.route('/results')
def results():    
    with open("text_files/edited.txt", "r", encoding='utf-8', errors="ignore") as f:
        edited_text = f.read()
    edited_text = edited_text.split("\n")
    return render_template("results.html", text_to_display=edited_text)


@app.route('/download')
def download(): 
    file_type = request.args.get('type')
    if file_type == "txt":
        return send_file("text_files/edited.txt", as_attachment=True)
    
    #*Would be nice to match the original document's formatting. Otherwise it is difficult to reject changes. 
    if file_type == "docx":
        edited = docx.Document()
        with open("text_files/edited.txt", "r", encoding='utf-8', errors="ignore") as f:
            edited_text = f.read()
        edited_text = edited_text.split("\n")
        for paragraph in edited_text:
            edited.add_paragraph(paragraph)
        edited.save("text_files/edited.docx")
        return send_file("text_files/edited.docx")


if __name__ == "__main__":
    app.run(debug=True)
